"""Unit tests for DOCX resume parsing."""
from __future__ import annotations

import io

import docx
import pytest

from app.gmail.client import Attachment
from app.pipeline.resume import (
    extract_docx_link_uris,
    extract_docx_text,
    parse_resume,
)


def _build_docx_with_hyperlink(text: str, url: str, anchor: str) -> bytes:
    """Build a minimal .docx with a paragraph, a table, and a real hyperlink."""
    doc = docx.Document()
    doc.add_paragraph(text)

    # Hyperlink — python-docx doesn't expose a high-level API, so we add the
    # relationship directly and append a hyperlink run to a paragraph.
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement

    part = doc.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    p = doc.add_paragraph()
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = anchor
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skill"
    table.rows[0].cells[1].text = "Python"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def docx_bytes() -> bytes:
    return _build_docx_with_hyperlink(
        "Jane Doe — Senior Engineer\n5 years experience",
        "https://linkedin.com/in/janedoe",
        "LinkedIn",
    )


def test_extract_docx_text_includes_paragraphs_and_tables(docx_bytes: bytes):
    text = extract_docx_text(docx_bytes)
    assert "Jane Doe" in text
    assert "Senior Engineer" in text
    assert "Python" in text  # from table cell


def test_extract_docx_link_uris_finds_hyperlinks(docx_bytes: bytes):
    urls = extract_docx_link_uris(docx_bytes)
    assert "https://linkedin.com/in/janedoe" in urls


def test_extract_docx_text_handles_empty_and_garbage():
    assert extract_docx_text(b"") == ""
    assert extract_docx_text(b"not a docx") == ""
    assert extract_docx_link_uris(b"") == []
    assert extract_docx_link_uris(b"not a docx") == []


def test_parse_resume_selects_docx(docx_bytes: bytes):
    att = Attachment(
        filename="resume.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=docx_bytes,
    )
    parsed = parse_resume([att])
    assert parsed.resume_present is True
    assert parsed.selected_filename == "resume.docx"
    assert "Jane Doe" in parsed.text
    assert "https://linkedin.com/in/janedoe" in parsed.urls


def test_parse_resume_prefers_pdf_over_docx(docx_bytes: bytes):
    # Minimal valid PDF bytes so PyMuPDF can open it.
    import fitz
    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    page.insert_text((72, 72), "Hello from PDF")
    pdf_bytes = pdf_doc.tobytes()
    pdf_doc.close()

    pdf_att = Attachment(filename="r.pdf", mime_type="application/pdf", data=pdf_bytes)
    docx_att = Attachment(
        filename="r.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=docx_bytes,
    )
    parsed = parse_resume([docx_att, pdf_att])
    assert parsed.selected_filename == "r.pdf"
    assert "Hello from PDF" in parsed.text
