"""Resume parsing.

Inputs: list of attachments from a Gmail message.
Outputs:
  - selected_resume: the PDF or DOCX attachment we treat as the resume (or None)
  - resume_text: extracted full text
  - urls_in_resume: URLs discovered in the resume
  - resume_present: True if a supported resume (PDF or DOCX) was found
"""
from __future__ import annotations

import io
from dataclasses import dataclass

import fitz  # PyMuPDF

from app.gmail.client import Attachment
from app.pipeline.extract import find_urls


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@dataclass
class ParsedResume:
    selected_filename: str | None
    text: str
    urls: list[str]
    resume_present: bool
    any_attachment: bool
    # Parse statistics for logging
    file_format: str | None = None  # "pdf" | "docx" | None
    text_length: int = 0
    url_count_from_text: int = 0
    url_count_from_annotations: int = 0
    parse_errors: list[str] | None = None


def _is_pdf(a: Attachment) -> bool:
    return (a.mime_type or "").lower() == "application/pdf" or a.filename.lower().endswith(".pdf")


def _is_docx(a: Attachment) -> bool:
    return (a.mime_type or "").lower() == DOCX_MIME or a.filename.lower().endswith(".docx")


def parse_resume(attachments: list[Attachment]) -> ParsedResume:
    any_attachment = bool(attachments)
    errors: list[str] = []

    # Prefer PDF, fall back to DOCX.
    pdf = next((a for a in attachments if _is_pdf(a)), None)
    if pdf is not None:
        text, urls, text_url_count, annot_url_count, errs = _parse_pdf_with_stats(pdf.data)
        errors.extend(errs)
        return ParsedResume(
            selected_filename=pdf.filename,
            text=text,
            urls=urls,
            resume_present=True,
            any_attachment=any_attachment,
            file_format="pdf",
            text_length=len(text),
            url_count_from_text=text_url_count,
            url_count_from_annotations=annot_url_count,
            parse_errors=errors or None,
        )

    docx_att = next((a for a in attachments if _is_docx(a)), None)
    if docx_att is not None:
        text, urls, text_url_count, annot_url_count, errs = _parse_docx_with_stats(docx_att.data)
        errors.extend(errs)
        return ParsedResume(
            selected_filename=docx_att.filename,
            text=text,
            urls=urls,
            resume_present=True,
            any_attachment=any_attachment,
            file_format="docx",
            text_length=len(text),
            url_count_from_text=text_url_count,
            url_count_from_annotations=annot_url_count,
            parse_errors=errors or None,
        )

    return ParsedResume(None, "", [], resume_present=False, any_attachment=any_attachment)


def extract_pdf_text(data: bytes) -> str:
    if not data:
        return ""
    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            parts = []
            for page in doc:
                parts.append(page.get_text("text"))
            return "\n".join(parts)
    except Exception:
        return ""


def extract_pdf_link_uris(data: bytes) -> list[str]:
    """Extract URIs from PDF link annotations.

    PDFs often hyperlink visible words like "LinkedIn" or "GitHub" — those URIs
    live in link annotations, not in the rendered text stream, so a plain
    text-extract pass misses them. This walks every page's link list and
    collects any `uri` entries.
    """
    if not data:
        return []
    out: list[str] = []
    seen: set[str] = set()
    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page in doc:
                for link in page.get_links():
                    uri = link.get("uri")
                    if uri and uri not in seen:
                        seen.add(uri)
                        out.append(uri)
    except Exception:
        return out
    return out


def extract_docx_text(data: bytes) -> str:
    if not data:
        return ""
    try:
        import docx  # python-docx

        document = docx.Document(io.BytesIO(data))
        parts: list[str] = []
        for para in document.paragraphs:
            if para.text:
                parts.append(para.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def extract_docx_link_uris(data: bytes) -> list[str]:
    """Extract URIs from DOCX hyperlink relationships.

    DOCX stores hyperlinks (e.g. linked "LinkedIn"/"GitHub" text) as relationships
    in `document.part.rels` rather than inline in the rendered text, mirroring
    PDF link annotations. Walk the rels and collect External hyperlink targets.
    """
    if not data:
        return []
    out: list[str] = []
    seen: set[str] = set()
    try:
        import docx  # python-docx

        document = docx.Document(io.BytesIO(data))
        for rel in document.part.rels.values():
            reltype = getattr(rel, "reltype", "") or ""
            if reltype.endswith("/hyperlink"):
                target = getattr(rel, "target_ref", None) or getattr(rel, "target", None)
                if target and target not in seen:
                    seen.add(str(target))
                    out.append(str(target))
    except Exception:
        return out
    return out


def _merge_urls(text_urls: list[str], annot_urls: list[str]) -> list[str]:
    merged: list[str] = []
    seen: set[str] = set()
    for u in text_urls + annot_urls:
        if u not in seen:
            seen.add(u)
            merged.append(u)
    return merged


def _parse_pdf_with_stats(data: bytes) -> tuple[str, list[str], int, int, list[str]]:
    """Returns (text, merged_urls, text_url_count, annot_url_count, errors)."""
    errors: list[str] = []
    text = extract_pdf_text(data)
    if not text and data:
        errors.append("pdf_text_extraction_empty")
    text_urls = find_urls(text)
    annot_urls = extract_pdf_link_uris(data)
    return text, _merge_urls(text_urls, annot_urls), len(text_urls), len(annot_urls), errors


def _parse_docx_with_stats(data: bytes) -> tuple[str, list[str], int, int, list[str]]:
    """Returns (text, merged_urls, text_url_count, annot_url_count, errors)."""
    errors: list[str] = []
    text = extract_docx_text(data)
    if not text and data:
        errors.append("docx_text_extraction_empty")
    text_urls = find_urls(text)
    annot_urls = extract_docx_link_uris(data)
    return text, _merge_urls(text_urls, annot_urls), len(text_urls), len(annot_urls), errors


def parse_pdf_bytes(data: bytes) -> tuple[str, list[str]]:
    text = extract_pdf_text(data)
    return text, _merge_urls(find_urls(text), extract_pdf_link_uris(data))


def parse_docx_bytes(data: bytes) -> tuple[str, list[str]]:
    text = extract_docx_text(data)
    return text, _merge_urls(find_urls(text), extract_docx_link_uris(data))
